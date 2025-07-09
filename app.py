import os
import logging
from flask import Flask, render_template, request, jsonify
from werkzeug.utils import secure_filename
import PyPDF2
import docx
from ml_analyzer import ResumeAnalyzer

# Pre-built job descriptions for different roles
JOB_TEMPLATES = {
    "software_engineer": {
        "title": "Software Engineer",
        "description": """We are seeking a skilled Software Engineer to join our development team. The ideal candidate will have experience in designing, developing, and maintaining software applications.

Key Responsibilities:
• Design and develop scalable software solutions
• Write clean, maintainable, and efficient code
• Collaborate with cross-functional teams
• Participate in code reviews and testing
• Debug and troubleshoot issues
• Stay updated with latest technologies

Required Skills:
• Bachelor's degree in Computer Science or related field
• 3+ years of experience in software development
• Proficiency in Python, Java, or JavaScript
• Experience with databases (SQL, NoSQL)
• Knowledge of version control systems (Git)
• Understanding of software development lifecycle
• Strong problem-solving skills
• Excellent communication skills

Preferred Skills:
• Experience with cloud platforms (AWS, Azure)
• Knowledge of containerization (Docker, Kubernetes)
• Familiarity with Agile methodologies
• Experience with testing frameworks"""
    },
    "data_scientist": {
        "title": "Data Scientist",
        "description": """We are looking for a Data Scientist to analyze large amounts of raw information to find patterns that will help improve our company's performance.

Key Responsibilities:
• Collect, process, and analyze large datasets
• Build predictive models and machine learning algorithms
• Create data visualizations and reports
• Collaborate with business stakeholders
• Present findings to management teams
• Develop data-driven solutions

Required Skills:
• Master's degree in Data Science, Statistics, or related field
• 2+ years of experience in data analysis
• Proficiency in Python or R
• Experience with SQL and databases
• Knowledge of machine learning algorithms
• Strong statistical analysis skills
• Experience with data visualization tools (Tableau, PowerBI)

Preferred Skills:
• Experience with big data technologies (Spark, Hadoop)
• Knowledge of deep learning frameworks (TensorFlow, PyTorch)
• Cloud platform experience (AWS, GCP)
• Experience with A/B testing
• Knowledge of business intelligence tools"""
    },
    "frontend_developer": {
        "title": "Frontend Developer",
        "description": """We are seeking a talented Frontend Developer to create engaging and user-friendly web interfaces.

Key Responsibilities:
• Develop responsive web applications
• Implement user interface designs
• Optimize applications for maximum speed
• Collaborate with UX/UI designers
• Ensure cross-browser compatibility
• Write reusable and maintainable code

Required Skills:
• Bachelor's degree in Computer Science or related field
• 2+ years of frontend development experience
• Proficiency in HTML, CSS, and JavaScript
• Experience with React, Angular, or Vue.js
• Knowledge of responsive design principles
• Understanding of version control (Git)
• Experience with build tools (Webpack, Gulp)

Preferred Skills:
• Experience with TypeScript
• Knowledge of CSS preprocessors (Sass, Less)
• Familiarity with testing frameworks (Jest, Cypress)
• Understanding of SEO principles
• Experience with progressive web apps (PWA)
• Knowledge of accessibility standards"""
    },
    "product_manager": {
        "title": "Product Manager",
        "description": """We are looking for a Product Manager to guide the success of our products and lead cross-functional teams.

Key Responsibilities:
• Define product vision and strategy
• Conduct market research and competitive analysis
• Gather and prioritize product requirements
• Work closely with engineering teams
• Analyze product performance metrics
• Manage product roadmap and timeline

Required Skills:
• Bachelor's degree in Business, Engineering, or related field
• 3+ years of product management experience
• Strong analytical and problem-solving skills
• Excellent communication and leadership skills
• Experience with agile development methodologies
• Knowledge of product analytics tools
• Understanding of user experience principles

Preferred Skills:
• MBA or advanced degree
• Experience in technology products
• Knowledge of A/B testing and experimentation
• Familiarity with design thinking principles
• Experience with project management tools (Jira, Asana)
• Understanding of technical concepts"""
    },
    "marketing_specialist": {
        "title": "Digital Marketing Specialist",
        "description": """We are seeking a Digital Marketing Specialist to develop and execute marketing campaigns across various digital channels.

Key Responsibilities:
• Develop and implement digital marketing strategies
• Manage social media accounts and campaigns
• Create engaging content for various platforms
• Analyze campaign performance and ROI
• Conduct market research and competitor analysis
• Collaborate with creative teams

Required Skills:
• Bachelor's degree in Marketing, Communications, or related field
• 2+ years of digital marketing experience
• Proficiency in Google Analytics and Google Ads
• Experience with social media platforms
• Strong written and verbal communication skills
• Knowledge of SEO and SEM principles
• Creative thinking and analytical skills

Preferred Skills:
• Experience with marketing automation tools
• Knowledge of email marketing platforms
• Familiarity with graphic design software
• Understanding of conversion optimization
• Experience with influencer marketing
• Knowledge of content management systems"""
    },
    "devops_engineer": {
        "title": "DevOps Engineer",
        "description": """We are looking for a DevOps Engineer to help us build and maintain our infrastructure and deployment pipelines.

Key Responsibilities:
• Design and implement CI/CD pipelines
• Manage cloud infrastructure and services
• Monitor system performance and reliability
• Automate deployment and scaling processes
• Ensure security best practices
• Collaborate with development teams

Required Skills:
• Bachelor's degree in Computer Science or related field
• 3+ years of DevOps or systems administration experience
• Experience with cloud platforms (AWS, Azure, GCP)
• Proficiency in containerization (Docker, Kubernetes)
• Knowledge of infrastructure as code (Terraform, CloudFormation)
• Experience with CI/CD tools (Jenkins, GitLab CI)
• Strong scripting skills (Bash, Python)

Preferred Skills:
• Experience with monitoring tools (Prometheus, Grafana)
• Knowledge of configuration management (Ansible, Chef)
• Understanding of security practices
• Experience with microservices architecture
• Familiarity with log management systems
• Knowledge of database administration"""
    }
}

# Configure logging
logging.basicConfig(level=logging.DEBUG)

# Create the Flask app
app = Flask(__name__)
app.secret_key = os.environ.get("SESSION_SECRET", "dev-secret-key")

# Configure file uploads
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}
MAX_CONTENT_LENGTH = 16 * 1024 * 1024  # 16MB max file size

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# Create upload directory if it doesn't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Initialize the ML analyzer
analyzer = ResumeAnalyzer()

def allowed_file(filename):
    """Check if file extension is allowed."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_path):
    """Extract text from PDF file."""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
    except Exception as e:
        app.logger.error(f"Error extracting PDF text: {str(e)}")
        return None

def extract_text_from_docx(file_path):
    """Extract text from DOCX file."""
    try:
        doc = docx.Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        
        return "\n".join(text).strip()
    except Exception as e:
        app.logger.error(f"Error extracting DOCX text: {str(e)}")
        return None

def extract_text_from_file(file_path, filename):
    """Extract text from uploaded file based on its extension."""
    file_extension = filename.rsplit('.', 1)[1].lower()
    
    if file_extension == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension == 'docx':
        return extract_text_from_docx(file_path)
    elif file_extension == 'txt':
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            app.logger.error(f"Error reading text file: {str(e)}")
            return None
    
    return None

@app.route('/')
def index():
    """Main page with resume optimization form."""
    return render_template('index.html')

@app.route('/job-templates')
def get_job_templates():
    """API endpoint to get available job templates."""
    return jsonify({
        'templates': {
            key: {'title': template['title']} 
            for key, template in JOB_TEMPLATES.items()
        }
    })

@app.route('/job-templates/<template_id>')
def get_job_template(template_id):
    """API endpoint to get a specific job template."""
    if template_id in JOB_TEMPLATES:
        return jsonify(JOB_TEMPLATES[template_id])
    else:
        return jsonify({'error': 'Template not found'}), 404

@app.route('/analyze', methods=['POST'])
def analyze_resume():
    """Analyze resume against job description and return results."""
    try:
        resume_text = ""
        job_description = ""
        
        # Check if it's a file upload request
        if 'resume_file' in request.files:
            resume_file = request.files['resume_file']
            job_description = request.form.get('job_description', '').strip()
            
            if resume_file and resume_file.filename and resume_file.filename != '':
                if allowed_file(resume_file.filename):
                    # Save the uploaded file
                    filename = secure_filename(resume_file.filename or 'resume')
                    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    resume_file.save(file_path)
                    
                    # Extract text from the file
                    resume_text = extract_text_from_file(file_path, filename)
                    
                    # Clean up the uploaded file
                    try:
                        os.remove(file_path)
                    except:
                        pass
                    
                    if not resume_text:
                        return jsonify({'error': 'Could not extract text from the uploaded file. Please check the file format.'}), 400
                else:
                    return jsonify({'error': 'Invalid file format. Please upload a PDF, DOCX, or TXT file.'}), 400
            else:
                # Fall back to text input if no file uploaded
                resume_text = request.form.get('resume', '').strip()
        else:
            # Handle JSON request (text input)
            data = request.get_json()
            
            if not data:
                return jsonify({'error': 'No data provided'}), 400
            
            resume_text = data.get('resume', '').strip()
            job_description = data.get('job_description', '').strip()
        
        if not resume_text:
            return jsonify({'error': 'Resume text is required'}), 400
        
        if not job_description:
            return jsonify({'error': 'Job description is required'}), 400
        
        # Analyze the resume
        results = analyzer.analyze_resume(resume_text, job_description)
        
        return jsonify(results)
    
    except Exception as e:
        app.logger.error(f"Error analyzing resume: {str(e)}")
        return jsonify({'error': 'An error occurred while analyzing the resume. Please try again.'}), 500

@app.errorhandler(404)
def not_found(error):
    return render_template('index.html'), 404

@app.errorhandler(500)
def internal_error(error):
    return jsonify({'error': 'Internal server error occurred'}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
